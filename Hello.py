# Copyright (c) Streamlit Inc. (2018-2022) Snowflake Inc. (2022)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import streamlit as st
from streamlit.logger import get_logger
from docx import Document
from io import BytesIO

LOGGER = get_logger(__name__)

def process_fasta(fasta_content):
    lines = fasta_content.splitlines()
    new_content = []

    sequence = ""
    for line in lines:
        if line.startswith(">"):
            if sequence:
                new_content.append(sequence)
                sequence = ""
            new_content.append(line)
        else:
            sequence += line.strip()

    if sequence:
        new_content.append(sequence)

    return "\n".join(new_content)

def read_docx(file_stream):
    doc = Document(file_stream)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

st.title("Clear your fasta file!")

uploaded_file = st.file_uploader("Upload your file in docx format!", type=["docx"])

if uploaded_file is not None:
    docx_content = read_docx(uploaded_file)
    result = process_fasta(docx_content)

    file_format = st.radio("Choose your output format:", ["txt", "docx"])

    if file_format == "txt":
        if st.button('Download your prepared file in .txt'):
            b_stream = BytesIO()
            b_stream.write(result.encode())
            st.download_button(
                label="Download .txt",
                data=b_stream,
                file_name="clean_fasta.txt",
                mime="text/plain"
            )
    else:
        if st.button('Download your prepared file in .docx'):
            b_stream = BytesIO()
            new_doc = Document()
            new_doc.add_paragraph(result)
            new_doc.save(b_stream)
            st.download_button(
                label="Download .docx",
                data=b_stream,
                file_name="clean_fasta.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

else:
    st.markdown(' ')
    
